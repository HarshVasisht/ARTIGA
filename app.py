from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests
from rich import traceback
traceback.install()
from dotenv import load_dotenv

load_dotenv()
PEXEL_API_KEY = os.getenv('PEXEL_API_KEY')
#Loading Quantized Llama model

def load_llm(max_tokens, Prompt_template):
    llm = CTransformers(
        model = "llm/llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type = "llama",
        max_new_tokens = max_tokens,
        temperature = 0.1
    )
    llm_chain = LLMChain(
        llm = llm,
        prompt = PromptTemplate.from_template(Prompt_template)
    )
    print(llm_chain)
    return llm

def fetch_photo(query):
    api_key =  os.getenv('PEXEL_API_KEY')

    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': api_key,
    }

    params = {
        'query': query,
        'per_page': 1,   # change to add more photos.
    }

    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        # print(photos)
        if photos:
            src_original_url = photos[0]['src']['original']
            # img_list = [photos[0]['src']['original'], photos[1]['src']['original']]
            return src_original_url
        else:
            print("No photos found for the given query.")
    else:
        print(f"Error: {response.status_code}, {response.text}")
    
    return None

def create_docx(user_input, paragraph, img_input):
    doc = Document()

    doc.add_heading(user_input, level = 1) 
    doc.add_paragraph(paragraph)
    doc.add_heading('Image', level = 1)
    image_stream = io.BytesIO()
    img_input.save(image_stream,format = 'PNG')
    img_input.seek(0)
    doc.add_picture(image_stream, width = Inches(4))

    return doc

st.set_page_config(layout = 'wide')

def main():

    st.title("A.R.T.I.G.A: AI for Real-Time Image Generation & Article creation")
    app_banner= Image.open('docs\img\index3.png')
    st.image(app_banner, caption='Future of Content generation')
    user_input = st.text_input("Enter your Topic")
    img_input = st.text_input("Enter your Topic for image")

    if len(user_input) and len(img_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])

        with col1:
            st.subheader("Generated content by Llama 2")
            st.write("Topic of the article: "+ user_input)
            st.write("Topic of the image: "+ img_input)

            PromptTemplate = """You are a digital marketing and SEO expert and your task is to generate articles for a given topic. So, write an article on {user_input} of 800 words.
            Stick to the topic given by the user and maintain a professional and creative tone. You can use quotes to go with the article."""
        llm_call = load_llm(max_tokens = 1000, Prompt_template = PromptTemplate)
        print(llm_call)
        result = llm_call(user_input)

        if len(result)>0:
            st.info("Ring Ring 🔔 , Article delivery! ")
            st.write(result)
        else:
            st.error("Oops, time's up. Couldn't find any relevent data on the given topic. ",icon="🚨")
            
        # with col2:
        #     st.subheader("Your Article's image just arrived")
        #     image_url = "temp_image.jpg"
        #     st.image(image_url)

        with col2:
            st.subheader("Fetched Image")
            image_url = fetch_photo(img_input)
            st.image(image_url)

        with col3:
            st.subheader("Final Article to Download")
            image_response = requests.get(image_url)
            img = Image.open(io.BytesIO(image_response.content))
            doc = create_docx(user_input, result['text'], img)

            # Save the Word document to a BytesIO buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Prepare the download link
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
if __name__ == "__main__":
    main()