from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests
from rich import traceback
traceback.install()
from dotenv import load_dotenv

load_dotenv()
#Loading the model
def load_llm(max_tokens, prompt_template):
    # Load the locally downloaded model here
    llm = CTransformers(
        model = "llm\llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.7
    )
    
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    print(llm_chain)
    return llm_chain

def get_src_original_url(query):
    PEXEL_API_KEY = os.getenv('PEXEL_API_KEY')
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': PEXEL_API_KEY,
    }

    params = {
        'query': query,
        'per_page': 1,
    }

    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            st.write("No photos found for the given query.")
    else:
        st.write(f"Error: {response.status_code}, {response.text}")

    return None

def create_word_docx(user_input, paragraph, image_input):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    # Add the image to the document
    doc.add_heading('Image Input', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))  # Adjust the width as needed

    return doc

st.set_page_config(layout="wide")

def main():
    st.title("A.R.T.I.G.A: AI for Real-Time Image Generation & Article creation")
    app_banner= Image.open('docs\img\index3.png')
    st.image(app_banner, caption='Future of Content generation')
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")

    image_input = st.text_input("Please enter the topic for the image you want to fetch!")

    if len(user_input) > 0 and len(image_input) > 0:

        col1, col2, col3 = st.columns([1,2,1])

        with col1:
            st.subheader("Generated Content by Llama 2")
            st.write("Topic of the article is: " + user_input)
            st.write("Image of the article is: " + image_input)
            # prompt_template = """You are a digital marketing and SEO expert and your task is to write article so write an article on the given topic: {user_input}. The article must be under 800 words. The article should be be lengthy.               
            prompt_template = """You are a digital marketing and SEO expert and your task is to generate articles for a given topic. So, write an article on {user_input} under 200 words.
            Stick to the topic given by the user and maintain a professional and creative tone. You can use quotes to go with the article. Also when reaching the word limit make sure no sentence is left uncompleted."""
            llm_call = load_llm(max_tokens=800, prompt_template=prompt_template)
            print(llm_call)
            result = llm_call(user_input)
            if len(result) > 0:
                st.info("Your article has been been generated successfully!")
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")

        with col2:
            st.subheader("Fetched Image")
            image_url = get_src_original_url(image_input)
            st.image(image_url)

        with col3:
            st.subheader("Final Article to Download")
            # image_input = "temp_image.jpg"
            image_response = requests.get(image_url)
            img = Image.open(io.BytesIO(image_response.content))
            doc = create_word_docx(user_input, result['text'], img)

            # Save the Word document to a BytesIO buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Prepare the download link
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )


if __name__ == "__main__":
    main()